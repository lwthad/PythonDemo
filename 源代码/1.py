#读取.doc文档内容，进行词频统计
excludes={"an","a","of","is","to","for","and"
          ,"the","or","that","as","in","are"
          ,"be","this","may","one"}
import docx

import sys
import pickle
import re
import codecs
import string
import shutil
from win32com import client as wc
def doSaveAas(): #doc格式转换为docx
    word = wc.Dispatch('Word.Application')
    
    doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\43-44.doc')  # 目标路径下的文件
    doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\1.docx', 12, False, "", True, "", False, False, False, False) # 转化后路径下的文件 

    doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\45-46.doc')
    doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\2.docx', 12, False, "", True, "", False, False, False, False)

    for i in range(5): #2*i + 48 - 2*i + 49
        doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\{}-{}.doc'.format(2*i+48,2*i+49))
        doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\{}.docx'.format(i+3), 12, False, "", True, "", False, False, False, False)

    doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\58-60.doc')
    doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\8.docx', 12, False, "", True, "", False, False, False, False)

    doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\61-62.doc')
    doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\9.docx', 12, False, "", True, "", False, False, False, False)
    
    for i in range(4): #4*i + 63 - 4*i + 65
        doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\{}-{}.doc'.format(4*i+63,4*i+65))
        doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\{}.docx'.format(i+10), 12, False, "", True, "", False, False, False, False)
        
    doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\74-76.doc')
    doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\14.docx', 12, False, "", True, "", False, False, False, False)

    doc = word.Documents.Open(r'C:\Users\tao\Desktop\2班\En\78-80.doc')
    doc.SaveAs(r'C:\Users\tao\Desktop\2班\En\docx\15.docx', 12, False, "", True, "", False, False, False, False)
    
    doc.Close()
    word.Quit()
    
doSaveAas() #调用函数 进行格式转换

def readDocument():
    doc = ""
    for i in range(15):
        file = docx.Document(r'C:\Users\tao\Desktop\2班\En\docx\{}.docx'.format(i+1))
        for para in file.paragraphs:
            doc = doc + para.text
    doc = doc.lower()
    for ch in '!"#$%&()*+,-./:;<=>?@[\\]^_‘{|}~':
        doc = doc.replace(ch, " ")  
    return doc

txtfile = readDocument()
words  = txtfile.split() #进行分词

counts = {} #初始化字典
for word in words:			
    counts[word] = counts.get(word,0) + 1
    
for word in excludes:
    del counts[word]

items = list(counts.items())

items.sort(key=lambda x:x[1], reverse=True) 
for i in range(10):
    word, count = items[i]
    print ('{0:<15}{1:>5}'.format(word, count))

   
